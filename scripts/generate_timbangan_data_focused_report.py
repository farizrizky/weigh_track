from __future__ import annotations

from datetime import datetime
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

import generate_timbangan_analysis_report as base


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = ROOT / "reports"
ASSET_DIR = OUTPUT_DIR / "assets"
OUTPUT_DOCX = OUTPUT_DIR / "Laporan_Analisis_Berbasis_Data_Digital_Per_Lot_vs_RAM.docx"
RECONSTRUCTION_PATH = ASSET_DIR / "rekonstruksi_empat_pasangan_timbang.png"
TIMELINE_PATH = ASSET_DIR / "timeline_susut_juli_agustus.png"


FULL_PAIRS = [
    {
        "no": 1,
        "digital_date": "11 Agu 2026",
        "ram_date": "11 Agu 2026",
        "origin": "Gembung",
        "truck": "BD 8450 K",
        "digital_gross": 13982,
        "digital_tare": 4250,
        "digital_net": 9732,
        "ram_gross": 12380,
        "ram_tare": 4250,
        "ram_net": 8130,
        "timing": "Hari yang sama",
        "delay_days": 0,
    },
    {
        "no": 2,
        "digital_date": "13 Agu 2026",
        "ram_date": "14 Agu 2026",
        "origin": "Sebayur",
        "truck": "BK 8035 WS",
        "digital_gross": 14633,
        "digital_tare": 4720,
        "digital_net": 9913,
        "ram_gross": 13520,
        "ram_tare": 4720,
        "ram_net": 8800,
        "timing": "Hari berikutnya",
        "delay_days": 1,
    },
    {
        "no": 3,
        "digital_date": "24 Agu 2026",
        "ram_date": "24 Agu 2026",
        "origin": "Sebayur",
        "truck": "BD 8450 K",
        "digital_gross": 15340,
        "digital_tare": 4250,
        "digital_net": 11090,
        "ram_gross": 14550,
        "ram_tare": 4250,
        "ram_net": 10300,
        "timing": "Hari yang sama",
        "delay_days": 0,
    },
    {
        "no": 4,
        "digital_date": "24 Agu 2026",
        "ram_date": "25 Agu 2026",
        "origin": "Gembung",
        "truck": "BK 8035 WS",
        "digital_gross": 14633,
        "digital_tare": 4700,
        "digital_net": 9933,
        "ram_gross": 12530,
        "ram_tare": 4700,
        "ram_net": 7830,
        "timing": "Hari berikutnya",
        "delay_days": 1,
    },
]

for row in FULL_PAIRS:
    row["gap"] = row["digital_net"] - row["ram_net"]
    row["gap_pct"] = row["gap"] / row["digital_net"]


def font(size: int, bold=False):
    candidates = [
        Path(r"C:\Windows\Fonts\arialbd.ttf" if bold else r"C:\Windows\Fonts\arial.ttf"),
        Path(r"C:\Windows\Fonts\calibrib.ttf" if bold else r"C:\Windows\Fonts\calibri.ttf"),
    ]
    for candidate in candidates:
        if candidate.exists():
            return ImageFont.truetype(str(candidate), size=size)
    return ImageFont.load_default()


def draw_right(draw, xy, text, text_font, fill):
    x, y = xy
    box = draw.textbbox((0, 0), text, font=text_font)
    draw.text((x - (box[2] - box[0]), y), text, font=text_font, fill=fill)


def make_reconstruction_image() -> None:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    width, height = 1800, 2240
    image = Image.new("RGB", (width, height), "#F7F9F8")
    draw = ImageDraw.Draw(image)

    title_font = font(42, True)
    sub_font = font(23)
    head_font = font(25, True)
    body_font = font(22)
    body_bold = font(23, True)
    small_font = font(19)

    draw.text((90, 60), "Rekonstruksi Empat Pasangan Timbang", font=title_font, fill="#1F4E3D")
    draw.text(
        (90, 118),
        "Disusun kembali dari screenshot yang diberikan | satuan kilogram",
        font=sub_font,
        fill="#68716D",
    )

    left = 90
    right = width - 90
    digital_x0, digital_x1 = 210, 890
    ram_x0, ram_x1 = 920, 1450
    gap_x0, gap_x1 = 1475, right

    y = 190
    draw.rounded_rectangle((left, y, right, y + 76), radius=8, fill="#1F4E3D")
    draw.text((left + 25, y + 23), "NO.", font=head_font, fill="white")
    draw.text((digital_x0 + 20, y + 23), "TIMBANGAN DIGITAL KEBUN", font=head_font, fill="white")
    draw.text((ram_x0 + 20, y + 23), "TIMBANGAN RAM", font=head_font, fill="white")
    gap_header = "SELISIH"
    gap_header_box = draw.textbbox((0, 0), gap_header, font=head_font)
    draw.text(
        ((gap_x0 + gap_x1 - (gap_header_box[2] - gap_header_box[0])) / 2, y + 23),
        gap_header,
        font=head_font,
        fill="white",
    )

    block_height = 465
    for index, row in enumerate(FULL_PAIRS):
        top = y + 96 + index * block_height
        bottom = top + 430
        fill = "#FFFFFF" if index % 2 == 0 else "#EEF4F0"
        draw.rounded_rectangle((left, top, right, bottom), radius=8, fill=fill, outline="#D7DEDA", width=2)
        draw.text((left + 35, top + 24), str(row["no"]), font=body_bold, fill="#1F4E3D")

        def panel(x0, x1, date_value, gross, tare, net):
            label_x = x0 + 25
            value_x = x1 - 25
            lines = [
                ("Tanggal timbang", date_value, False),
                ("Produksi", row["origin"], False),
                ("Nopol", row["truck"], False),
                ("Gross", base.format_kg(gross), False),
                ("Tare", base.format_kg(tare), False),
                ("Net", base.format_kg(net), True),
            ]
            line_y = top + 28
            for label, value, bold in lines:
                if label == "Gross":
                    line_y += 24
                    draw.line((label_x, line_y - 12, x1 - 25, line_y - 12), fill="#C9D2CD", width=2)
                draw.text((label_x, line_y), label, font=body_bold if bold else body_font, fill="#252B29")
                draw_right(draw, (value_x, line_y), value, body_bold if bold else body_font, "#1F4E3D" if bold else "#252B29")
                line_y += 56

        panel(
            digital_x0,
            digital_x1,
            row["digital_date"],
            row["digital_gross"],
            row["digital_tare"],
            row["digital_net"],
        )
        panel(ram_x0, ram_x1, row["ram_date"], row["ram_gross"], row["ram_tare"], row["ram_net"])

        draw.line((digital_x1 + 15, top + 20, digital_x1 + 15, bottom - 20), fill="#D7DEDA", width=2)
        draw.line((ram_x1 + 10, top + 20, ram_x1 + 10, bottom - 20), fill="#D7DEDA", width=2)
        gap_text = base.format_kg(row["gap"])
        gap_pct = base.format_pct(row["gap_pct"])
        draw_right(draw, (gap_x1 - 20, top + 165), gap_text, font(29, True), "#A13A32")
        draw_right(draw, (gap_x1 - 20, top + 210), gap_pct, body_bold, "#A13A32")
        draw_right(draw, (gap_x1 - 20, top + 260), row["timing"], small_font, "#68716D")

    draw.text(
        (90, height - 94),
        "Catatan: angka gross digital adalah angka pada screenshot; definisi apakah pembacaan langsung atau hasil rekonstruksi perlu dikonfirmasi.",
        font=small_font,
        fill="#68716D",
    )
    image.save(RECONSTRUCTION_PATH, quality=95)


def make_timeline_chart(source_rows) -> None:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    width, height = 1700, 850
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)

    title_font = font(38, True)
    sub_font = font(21)
    axis_font = font(18)
    label_font = font(17, True)
    draw.text((90, 42), "Timeline Susut Gabungan Juli-Agustus 2026", font=title_font, fill="#1F4E3D")
    draw.text((90, 93), "Warna menunjukkan metode yang tercatat pada workbook", font=sub_font, fill="#68716D")

    left, top, right, bottom = 120, 160, 1625, 680
    max_y = 0.55
    for pct in [0, 0.1, 0.2, 0.3, 0.4, 0.5]:
        y = bottom - pct / max_y * (bottom - top)
        draw.line((left, y, right, y), fill="#E7EBE9", width=2)
        draw_right(draw, (left - 18, y - 10), f"{int(pct * 100)}%", axis_font, "#68716D")

    dates = [row["date"] for row in source_rows]
    min_date, max_date = min(dates), max(dates)
    span_days = (max_date - min_date).days

    points = []
    for row in source_rows:
        x = left + ((row["date"] - min_date).days / span_days) * (right - left)
        y = bottom - row["combined"] / max_y * (bottom - top)
        points.append((x, y, row))

    for (x1, y1, _), (x2, y2, _) in zip(points, points[1:]):
        draw.line((x1, y1, x2, y2), fill="#D1D8D4", width=3)

    for x, y, row in points:
        is_digital = row["scale"].lower() == "digital"
        color = "#C79A3B" if is_digital else "#3E755E"
        radius = 11 if is_digital else 8
        draw.ellipse((x - radius, y - radius, x + radius, y + radius), fill=color, outline="white", width=3)
        if is_digital or row["combined"] == min(r["combined"] for r in source_rows):
            label = base.format_pct(row["combined"], 1)
            box = draw.textbbox((0, 0), label, font=label_font)
            draw.rounded_rectangle(
                (x - (box[2] - box[0]) / 2 - 8, y - 47, x + (box[2] - box[0]) / 2 + 8, y - 18),
                radius=5,
                fill="#FFFFFF",
                outline=color,
                width=2,
            )
            draw.text((x - (box[2] - box[0]) / 2, y - 45), label, font=label_font, fill=color)

    tick_dates = [
        datetime(2026, 7, 1),
        datetime(2026, 7, 10),
        datetime(2026, 7, 21),
        datetime(2026, 8, 1),
        datetime(2026, 8, 10),
        datetime(2026, 8, 19),
        datetime(2026, 8, 24),
    ]
    for date_value in tick_dates:
        x = left + ((date_value - min_date).days / span_days) * (right - left)
        draw.line((x, bottom, x, bottom + 8), fill="#68716D", width=2)
        label = date_value.strftime("%d %b").replace("Aug", "Agu")
        box = draw.textbbox((0, 0), label, font=axis_font)
        draw.text((x - (box[2] - box[0]) / 2, bottom + 16), label, font=axis_font, fill="#68716D")

    legend_y = 770
    draw.ellipse((630, legend_y, 652, legend_y + 22), fill="#3E755E")
    draw.text((665, legend_y - 3), "RAM", font=sub_font, fill="#252B29")
    draw.ellipse((790, legend_y, 812, legend_y + 22), fill="#C79A3B")
    draw.text((825, legend_y - 3), "Digital", font=sub_font, fill="#252B29")
    image.save(TIMELINE_PATH, quality=95)


def add_cover(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(30)
    base.add_run(p, "PT JULANG PLANTATIONS", bold=True, size=11, color=base.GREEN)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(36)
    base.add_run(p, "LAPORAN ANALISIS\nBERBASIS DATA", bold=True, size=27, color=base.GREEN)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    base.add_run(p, "Perbandingan Timbangan Digital Per Lot dan Timbangan RAM", bold=True, size=16, color=base.CHARCOAL)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(22)
    base.add_run(
        p,
        "Fokus pada empat pasangan timbang dan data susut Juli-Agustus 2026",
        size=11.5,
        color=base.MID_GREY,
    )

    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    base.write_cell(
        table.cell(0, 0),
        "STATUS: ANALISIS SEMENTARA - PENYEBAB BELUM DAPAT DITETAPKAN SECARA TUNGGAL",
        bold=True,
        size=9,
        color=base.GREEN,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        fill=base.PALE_GREEN,
    )
    base.set_cell_border(
        table.cell(0, 0),
        top={"val": "single", "sz": 8, "color": base.GREEN},
        bottom={"val": "single", "sz": 8, "color": base.GREEN},
        left={"val": "single", "sz": 8, "color": base.GREEN},
        right={"val": "single", "sz": 8, "color": base.GREEN},
    )

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(90)
    base.add_run(p, "Sumber analisis", bold=True, size=8.5, color=base.MID_GREY)
    base.add_run(p, "\nScreenshot empat pasangan timbang dan workbook Ekstrak Susut", bold=True, size=11, color=base.CHARCOAL)
    p = doc.add_paragraph()
    base.add_run(p, "Tanggal penyusunan", bold=True, size=8.5, color=base.MID_GREY)
    base.add_run(p, "\n26 Agustus 2026", bold=True, size=11, color=base.CHARCOAL)
    p = doc.add_paragraph()
    base.add_run(p, "Penggunaan", bold=True, size=8.5, color=base.MID_GREY)
    base.add_run(p, "\nInternal PT Julang Plantations", bold=True, size=11, color=base.CHARCOAL)
    doc.add_page_break()


def add_summary_cards(doc: Document) -> None:
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    base.set_table_fixed_layout(table)
    cards = [
        ("4 DO", "Pasangan data"),
        ("40.668 kg", "Total net digital"),
        ("35.060 kg", "Total net RAM"),
        ("5.608 kg | 13,79%", "Gap kumulatif"),
    ]
    for index, (value, label) in enumerate(cards):
        cell = table.cell(0, index)
        base.set_cell_shading(cell, base.PALE_GOLD if index == 3 else base.VERY_LIGHT_GREY)
        base.clear_cell(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        base.add_run(p, value, bold=True, size=12.5 if index != 3 else 11.5, color=base.GOLD if index == 3 else base.GREEN)
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(0)
        base.add_run(p2, label, size=8.2, color=base.MID_GREY)
        base.set_cell_margins(cell, top=170, start=80, bottom=170, end=80)
    base.set_table_borders(table, base.WHITE, 4)


def add_derived_pair_table(doc: Document) -> None:
    headers = ["DO", "Produksi", "Nopol", "Net digital", "Net RAM", "Gap", "Gap %", "Jeda tanggal"]
    widths = [0.7, 1.8, 2.0, 2.0, 1.9, 1.8, 1.5, 2.3]
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    base.set_table_fixed_layout(table)
    for index, header in enumerate(headers):
        base.write_cell(
            table.cell(0, index),
            header,
            bold=True,
            size=7.8,
            color=base.WHITE,
            align=WD_ALIGN_PARAGRAPH.CENTER,
            fill=base.GREEN,
        )
        base.set_cell_width(table.cell(0, index), widths[index])
    base.set_repeat_table_header(table.rows[0])
    for row_index, row in enumerate(FULL_PAIRS, start=1):
        values = [
            row["no"],
            row["origin"],
            row["truck"],
            f"{base.format_kg(row['digital_net'])} kg",
            f"{base.format_kg(row['ram_net'])} kg",
            f"{base.format_kg(row['gap'])} kg",
            base.format_pct(row["gap_pct"]),
            row["timing"],
        ]
        cells = table.add_row().cells
        fill = base.WHITE if row_index % 2 else base.VERY_LIGHT_GREY
        for index, value in enumerate(values):
            align = WD_ALIGN_PARAGRAPH.RIGHT if 3 <= index <= 6 else WD_ALIGN_PARAGRAPH.CENTER
            base.write_cell(cells[index], value, size=7.9, align=align, fill=fill)
            base.set_cell_width(cells[index], widths[index])
        base.write_cell(cells[5], values[5], bold=True, size=7.9, color=base.RED, align=WD_ALIGN_PARAGRAPH.RIGHT, fill=fill)
        base.write_cell(cells[6], values[6], bold=True, size=7.9, color=base.RED, align=WD_ALIGN_PARAGRAPH.RIGHT, fill=fill)
        base.prevent_row_split(table.rows[-1])
    base.set_table_borders(table, base.LIGHT_GREY, 5)


def add_gross_tare_observations(doc: Document) -> None:
    rows = [
        ("Tare sama dalam setiap pasangan", "Gap net pada tabel sama persis dengan gap gross. Rumus net = gross - tare bukan sumber selisih matematis."),
        ("Digital selalu lebih tinggi", "Menunjukkan penyebab yang arahnya sistematis pada rantai proses, bukan variasi acak yang berganti tanda."),
        ("Gross digital 14.633 kg muncul dua kali", "Perlu dikonfirmasi apakah pembacaan langsung, hasil penjumlahan net lot + tare, atau input manual."),
        ("Resolusi berbeda", "Gross RAM tampak dicatat per 10 kg, sedangkan digital per 1 kg. Pembulatan ini terlalu kecil untuk menjelaskan gap 790-2.103 kg."),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    base.set_table_fixed_layout(table)
    for index, header in enumerate(("Observasi data", "Makna")):
        base.write_cell(table.cell(0, index), header, bold=True, size=8.5, color=base.WHITE, align=WD_ALIGN_PARAGRAPH.CENTER, fill=base.GREEN)
    for row_index, values in enumerate(rows, start=1):
        cells = table.add_row().cells
        fill = base.WHITE if row_index % 2 else base.VERY_LIGHT_GREY
        base.write_cell(cells[0], values[0], bold=True, size=8.5, color=base.GREEN_2, fill=fill, valign=WD_CELL_VERTICAL_ALIGNMENT.TOP)
        base.write_cell(cells[1], values[1], size=8.5, fill=fill, valign=WD_CELL_VERTICAL_ALIGNMENT.TOP)
        base.prevent_row_split(table.rows[-1])
    base.set_table_borders(table, base.LIGHT_GREY, 5)


def add_process_comparison(doc: Document) -> None:
    table = doc.add_table(rows=3, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    base.set_table_fixed_layout(table)
    headers = ["Aspek", "Digital per lot", "RAM"]
    for index, header in enumerate(headers):
        base.write_cell(table.cell(0, index), header, bold=True, size=8.5, color=base.WHITE, align=WD_ALIGN_PARAGRAPH.CENTER, fill=base.GREEN)
    values = [
        ("Objek yang ditimbang", "Cup lump ditumpuk langsung pada platform, tanpa wadah", "Truk beserta seluruh muatan yang tersisa di dalamnya"),
        ("Titik waktu", "Saat masing-masing lot selesai disusun dan dibaca", "Setelah pemindahan, loading tambahan 0,5-1 jam, menunggu, dan perjalanan"),
    ]
    for row_index, row_values in enumerate(values, start=1):
        fill = base.WHITE if row_index % 2 else base.VERY_LIGHT_GREY
        for column, value in enumerate(row_values):
            base.write_cell(
                table.cell(row_index, column),
                value,
                bold=column == 0,
                size=8.7,
                color=base.GREEN_2 if column == 0 else base.CHARCOAL,
                fill=fill,
                valign=WD_CELL_VERTICAL_ALIGNMENT.TOP,
            )
    base.set_table_borders(table, base.LIGHT_GREY, 5)


def add_magnitude_table(doc: Document) -> None:
    rows = [
        ("DO 3 - Sebayur", "7,12%", "Terendah dalam sampel, tetapi tetap 790 kg."),
        ("DO 2 - Sebayur", "11,23%", "Lebih tinggi dan ditimbang RAM pada hari berikutnya."),
        ("DO 1 - Gembung", "16,46%", "Terjadi pada tanggal yang sama; waktu menginap tidak dapat menjadi penjelasan tunggal."),
        ("DO 4 - Gembung", "21,17%", "Tertinggi; kombinasi asal material, waktu, kendaraan, dan proses perlu diperiksa."),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    base.set_table_fixed_layout(table)
    headers = ["Pasangan", "Gap", "Pembacaan sementara"]
    widths = [3.4, 2.0, 10.0]
    for index, header in enumerate(headers):
        base.write_cell(table.cell(0, index), header, bold=True, size=8.5, color=base.WHITE, align=WD_ALIGN_PARAGRAPH.CENTER, fill=base.GREEN)
        base.set_cell_width(table.cell(0, index), widths[index])
    for row_index, values in enumerate(rows, start=1):
        fill = base.WHITE if row_index % 2 else base.VERY_LIGHT_GREY
        cells = table.add_row().cells
        for index, value in enumerate(values):
            base.write_cell(
                cells[index],
                value,
                bold=index in (0, 1),
                size=8.5,
                color=base.RED if index == 1 else (base.GREEN_2 if index == 0 else base.CHARCOAL),
                align=WD_ALIGN_PARAGRAPH.CENTER if index == 1 else WD_ALIGN_PARAGRAPH.LEFT,
                fill=fill,
                valign=WD_CELL_VERTICAL_ALIGNMENT.TOP,
            )
            base.set_cell_width(cells[index], widths[index])
        base.prevent_row_split(table.rows[-1])
    base.set_table_borders(table, base.LIGHT_GREY, 5)


def add_pattern_tables(doc: Document) -> None:
    headers = ["Kelompok", "Net digital", "Net RAM", "Gap", "Gap terhadap digital"]
    rows = [
        ("Hari yang sama (2 DO)", 20822, 18430, 2392, 2392 / 20822),
        ("Hari berikutnya (2 DO)", 19846, 16630, 3216, 3216 / 19846),
        ("Gembung (2 DO)", 19665, 15960, 3705, 3705 / 19665),
        ("Sebayur (2 DO)", 21003, 19100, 1903, 1903 / 21003),
    ]
    table = doc.add_table(rows=1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    base.set_table_fixed_layout(table)
    for index, header in enumerate(headers):
        base.write_cell(table.cell(0, index), header, bold=True, size=8.2, color=base.WHITE, align=WD_ALIGN_PARAGRAPH.CENTER, fill=base.GREEN)
    for row_index, row in enumerate(rows, start=1):
        fill = base.WHITE if row_index % 2 else base.VERY_LIGHT_GREY
        cells = table.add_row().cells
        values = [row[0], f"{base.format_kg(row[1])} kg", f"{base.format_kg(row[2])} kg", f"{base.format_kg(row[3])} kg", base.format_pct(row[4])]
        for index, value in enumerate(values):
            align = WD_ALIGN_PARAGRAPH.LEFT if index == 0 else WD_ALIGN_PARAGRAPH.RIGHT
            base.write_cell(cells[index], value, bold=index in (0, 4), size=8.5, color=base.GREEN_2 if index == 0 else (base.RED if index == 4 else base.CHARCOAL), align=align, fill=fill)
        base.prevent_row_split(table.rows[-1])
    base.set_table_borders(table, base.LIGHT_GREY, 5)


def add_relationship_table(doc: Document) -> None:
    rows = [
        ("Net digital vs net RAM", "+0,93", "Bergerak searah: muatan yang lebih berat di digital umumnya tetap lebih berat di RAM."),
        ("Net digital vs gap kg", "-0,69", "Tidak mendukung dugaan sederhana bahwa semakin besar total muatan maka gap semakin besar."),
        ("Net digital vs gap %", "-0,73", "Muatan terbesar justru memiliki persentase gap terendah."),
        ("Jeda tanggal vs gap", "+0,41", "Arahnya mendukung pengaruh waktu, tetapi seluruh hari berikutnya memakai BK 8035 WS."),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    base.set_table_fixed_layout(table)
    for index, header in enumerate(("Hubungan deskriptif", "Korelasi", "Interpretasi")):
        base.write_cell(table.cell(0, index), header, bold=True, size=8.2, color=base.WHITE, align=WD_ALIGN_PARAGRAPH.CENTER, fill=base.GREEN)
    for row_index, values in enumerate(rows, start=1):
        fill = base.WHITE if row_index % 2 else base.VERY_LIGHT_GREY
        cells = table.add_row().cells
        base.write_cell(cells[0], values[0], bold=True, size=8.3, color=base.GREEN_2, fill=fill)
        base.write_cell(cells[1], values[1], bold=True, size=9, color=base.RED, align=WD_ALIGN_PARAGRAPH.CENTER, fill=fill)
        base.write_cell(cells[2], values[2], size=8.3, fill=fill)
        base.prevent_row_split(table.rows[-1])
    base.set_table_borders(table, base.LIGHT_GREY, 5)


def add_factor_matrix(doc: Document) -> None:
    headers = ["Faktor", "Bagaimana menghasilkan gap", "Kesesuaian dengan data", "Status kesimpulan"]
    widths = [4.2, 7.0, 7.0, 5.0]
    rows = [
        (
            "Kadar air, KKK, umur, dan ukuran cup lump",
            "Digital menangkap berat material pada kondisi lebih awal; material yang lebih basah berpotensi kehilangan bobot lebih besar.",
            "Gembung 18,84% vs Sebayur 9,06%; literatur menunjukkan susut tinggi pada 1-2 hari awal dan dipengaruhi KKK.",
            "Indikasi kuat, belum diukur langsung",
        ),
        (
            "Air tertinggal di platform atau area pemindahan",
            "Air masih memberi beban saat digital dibaca tetapi tidak ikut masuk ke truk.",
            "Cup lump ditumpuk langsung tanpa wadah; volume air/residu belum dicatat.",
            "Sangat mungkin, perlu neraca air",
        ),
        (
            "Waktu, tekanan tumpukan, dan perjalanan",
            "Setelah digital, material disusun ulang, menunggu, dan terguncang sehingga air/material dapat keluar dari kendaraan.",
            "Hari berikutnya 16,20% vs hari sama 11,49%, tetapi kendaraan berbeda.",
            "Indikasi arah, faktor tercampur",
        ),
        (
            "Kesesuaian lot tertimbang dan lot termuat",
            "Lot tertinggal, berpindah DO, tercatat dua kali, atau tidak seluruhnya masuk truk.",
            "Belum ada daftar rekonsiliasi ID lot dan checklist fisik loading.",
            "Belum dapat dinilai",
        ),
        (
            "Zero, residu, dan stabilitas timbangan digital",
            "Air/potongan tertinggal, zero tidak konsisten, pembacaan belum stabil, atau beban tidak merata menimbulkan bias.",
            "Arah digital selalu lebih tinggi; log zero dan uji repeatability belum tersedia.",
            "Mungkin, perlu uji alat dan prosedur",
        ),
        (
            "Kendaraan, bak, dan rute",
            "Luas bak, celah drainase, kemiringan, getaran, dan kebocoran menentukan material/air yang tetap berada di truk.",
            "Hari sama selalu BD 8450 K; hari berikutnya selalu BK 8035 WS.",
            "Tidak terpisah dari efek waktu",
        ),
        (
            "Tare RAM dan konfigurasi truk",
            "Stored tare yang tidak aktual dapat membuat net RAM terlalu rendah atau terlalu tinggi.",
            "Tare sama dalam setiap pasangan, tetapi asal dan waktu penetapan tare belum diketahui.",
            "Mungkin, bukan penjelasan tunggal",
        ),
        (
            "Operasi dan kondisi timbangan RAM",
            "Posisi kendaraan, roda tidak seluruhnya di platform, binding, zero, rem, serta kalibrasi memengaruhi gross.",
            "Tiket, prosedur, dan bukti verifikasi RAM belum tersedia.",
            "Mungkin, perlu pemeriksaan independen",
        ),
        (
            "Kehilangan fisik atau perubahan muatan",
            "Cup lump jatuh, tertinggal, dipindahkan, atau air/material keluar setelah loading.",
            "Belum tersedia segel, foto, atau checklist sebelum-sesudah perjalanan.",
            "Belum dapat dinilai",
        ),
        (
            "Pembulatan dan resolusi",
            "Perbedaan skala pembacaan menghasilkan selisih kecil pada angka akhir.",
            "Gap 790-2.103 kg jauh lebih besar daripada pembulatan 1-10 kg.",
            "Tidak memadai sebagai penyebab utama",
        ),
    ]
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    base.set_table_fixed_layout(table)
    for index, header in enumerate(headers):
        base.write_cell(table.cell(0, index), header, bold=True, size=8, color=base.WHITE, align=WD_ALIGN_PARAGRAPH.CENTER, fill=base.GREEN)
        base.set_cell_width(table.cell(0, index), widths[index])
    base.set_repeat_table_header(table.rows[0])
    for row_index, values in enumerate(rows, start=1):
        fill = base.WHITE if row_index % 2 else base.VERY_LIGHT_GREY
        cells = table.add_row().cells
        for index, value in enumerate(values):
            base.write_cell(
                cells[index],
                value,
                bold=index in (0, 3),
                size=8,
                color=base.GREEN_2 if index == 0 else (base.RED if index == 3 else base.CHARCOAL),
                fill=fill,
                valign=WD_CELL_VERTICAL_ALIGNMENT.TOP,
            )
            base.set_cell_width(cells[index], widths[index])
        base.prevent_row_split(table.rows[-1])
    base.set_table_borders(table, base.LIGHT_GREY, 5)


def add_evidence_conclusions(doc: Document) -> None:
    base.add_callout(
        doc,
        "Fakta yang dapat dinyatakan",
        "Empat pasangan seluruhnya menunjukkan digital lebih tinggi; gap kumulatif 5.608 kg atau 13,79%; tare identik dalam tiap pasangan; variasi gap lebar; dan perbandingan dilakukan pada objek, waktu, serta kondisi yang tidak sama.",
        fill=base.PALE_GREEN,
        accent=base.GREEN,
    )
    base.add_callout(
        doc,
        "Indikasi yang paling kuat",
        "Kondisi material/asal produksi dan jeda proses kemungkinan berkontribusi. Gembung konsisten lebih tinggi daripada Sebayur, sementara kelompok hari berikutnya lebih tinggi daripada hari yang sama. Efek waktu tetap tercampur dengan kendaraan.",
        fill=base.PALE_GOLD,
        accent=base.GOLD,
    )
    base.add_callout(
        doc,
        "Yang belum boleh disimpulkan",
        "Belum dapat dinyatakan bahwa tekanan, kadar air, timbangan digital, timbangan RAM, tare, kendaraan, atau kesalahan transaksi merupakan penyebab tunggal. Data juga tidak mendukung klaim bahwa semakin besar total net digital maka gap otomatis semakin besar.",
        fill=base.PALE_RED,
        accent=base.RED,
    )


def add_required_data_table(doc: Document) -> None:
    headers = ["Data tambahan", "Mengapa penting", "Pertanyaan yang dijawab"]
    rows = [
        ("Timestamp digital pertama/terakhir, selesai loading, berangkat, dan RAM", "Mengubah kategori hari menjadi jeda aktual dalam jam", "Apakah gap meningkat seiring waktu?"),
        ("Usia sejak koagulasi, KKK/kadar air, ukuran, hujan, dan produksi", "Mengontrol kondisi material saat ditimbang", "Apakah Gembung memang lebih basah atau lebih muda?"),
        ("ID dan net setiap lot serta checklist lot naik truk", "Merekonsiliasi jumlah digital dengan muatan fisik", "Apakah seluruh lot digital benar-benar dikirim?"),
        ("Berat/volume air dan residu yang tertinggal di platform", "Membentuk neraca massa setelah penimbangan digital", "Berapa bagian gap yang sudah hilang sebelum truk berangkat?"),
        ("Luas bak, tinggi tumpukan, kondisi drainase, dan rute", "Memisahkan efek kendaraan dan susunan muatan", "Apakah BK mengeluarkan air lebih banyak daripada BD?"),
        ("Tare kosong aktual, tiket RAM, zero, dan bukti kalibrasi", "Memastikan net RAM dan gross kedua alat dapat dipercaya", "Apakah gap berasal dari alat/tare atau material?"),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    base.set_table_fixed_layout(table)
    widths = [6.0, 8.0, 6.0]
    for index, header in enumerate(headers):
        base.write_cell(table.cell(0, index), header, bold=True, size=8.2, color=base.WHITE, align=WD_ALIGN_PARAGRAPH.CENTER, fill=base.GREEN)
        base.set_cell_width(table.cell(0, index), widths[index])
    for row_index, values in enumerate(rows, start=1):
        fill = base.WHITE if row_index % 2 else base.VERY_LIGHT_GREY
        cells = table.add_row().cells
        for index, value in enumerate(values):
            base.write_cell(cells[index], value, bold=index == 0, size=8.2, color=base.GREEN_2 if index == 0 else base.CHARCOAL, fill=fill, valign=WD_CELL_VERTICAL_ALIGNMENT.TOP)
            base.set_cell_width(cells[index], widths[index])
        base.prevent_row_split(table.rows[-1])
    base.set_table_borders(table, base.LIGHT_GREY, 5)


def create_document() -> Path:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    source_rows = base.load_source_rows()
    make_reconstruction_image()
    make_timeline_chart(source_rows)
    base.make_chart()

    doc = Document()
    base.setup_styles(doc)
    base.configure_section(doc.sections[0])
    base.configure_header_footer(doc.sections[0])
    add_cover(doc)

    base.add_section_title(doc, "00", "Pertanyaan dan Batas Analisis")
    base.add_body_paragraph(
        doc,
        "Laporan ini menjawab dua pertanyaan: pertama, mengapa perbandingan digital per lot dan RAM hampir pasti menghasilkan selisih; kedua, apakah rentang selisih 7,12%-21,17% dapat dianggap wajar berdasarkan data yang tersedia.",
    )
    control = doc.add_table(rows=5, cols=2)
    control.alignment = WD_TABLE_ALIGNMENT.CENTER
    control_rows = [
        ("Sumber 1", "Screenshot empat pasangan gross, tare, dan net digital-RAM"),
        ("Sumber 2", "Workbook 'perbandingan susut ram dan digital.xlsx', sheet Ekstrak Susut"),
        ("Fakta proses", "Digital menimbang cup lump secara langsung di platform tanpa wadah"),
        ("Perubahan waktu", "Proses digital menambah sekitar 0,5-1 jam sebelum perjalanan ke RAM"),
        ("Batasan", "Hanya empat pasangan; tidak ada timestamp jam, KKK/kadar air, jumlah lot, neraca air, tare aktual, dan tiket kalibrasi"),
    ]
    for index, (label, value) in enumerate(control_rows):
        base.write_cell(control.cell(index, 0), label, bold=True, size=8.7, color=base.GREEN, fill=base.PALE_GREEN)
        base.write_cell(control.cell(index, 1), value, size=8.7, fill=base.WHITE)
    base.set_table_borders(control, base.LIGHT_GREY, 5)
    base.add_callout(
        doc,
        "Catatan workbook",
        "Seluruh baris Juli diberi label RAM. Laporan mengikuti label workbook dan menganggap maksud penjelasan awal adalah Juli belum menggunakan digital. Persentase dalam workbook berupa nilai jadi tanpa kuantitas dasar atau formula yang dapat diaudit.",
        fill=base.PALE_GOLD,
        accent=base.GOLD,
    )

    base.add_section_title(doc, "01", "Ringkasan Eksekutif")
    add_summary_cards(doc)
    base.add_callout(
        doc,
        "Jawaban utama",
        "Secara operasional, gap hampir pasti muncul karena digital dan RAM tidak menimbang objek pada waktu serta kondisi yang identik. Namun, data belum cukup untuk menyatakan bahwa gap sampai 21,17% merupakan susut yang wajar. Besarnya selisih harus diperlakukan sebagai hasil gabungan perubahan material, waktu, pemindahan/loading, kendaraan, rekonsiliasi lot, tare, serta ketidakpastian alat dan prosedur.",
    )
    for text in [
        "Tidak ada dasar untuk menetapkan satu penyebab sebagai satu-satunya sumber perbedaan.",
        "Total net digital tidak menunjukkan hubungan positif dengan besarnya gap; muatan terbesar justru mempunyai gap terendah.",
        "Jeda ke hari berikutnya menunjukkan gap lebih besar, tetapi seluruh kasus tersebut memakai kendaraan yang sama sehingga efek waktu dan kendaraan tercampur.",
        "Perbedaan Gembung dan Sebayur lebih besar daripada perbedaan kelompok waktu dan layak menjadi fokus pemeriksaan kondisi material.",
    ]:
        base.add_bullet(doc, text)

    base.add_section_title(doc, "02", "Rekonstruksi Screenshot Empat Pasangan")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(RECONSTRUCTION_PATH), width=Cm(16.3))
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    base.add_run(caption, "Gambar 1. Rekonstruksi visual data yang terdapat pada screenshot.", size=8, color=base.MID_GREY, italic=True)
    add_derived_pair_table(doc)
    add_gross_tare_observations(doc)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(base.CHART_PATH), width=Cm(16.3))
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    base.add_run(caption, "Gambar 2. Net digital lebih tinggi pada seluruh pasangan, tetapi besar gap tidak seragam.", size=8, color=base.MID_GREY, italic=True)

    base.add_section_title(doc, "03", "Mengapa Gap Hampir Pasti Terjadi")
    add_process_comparison(doc)
    base.add_body_paragraph(
        doc,
        "Perbandingan ini bukan uji dua timbangan terhadap objek yang sama pada waktu yang sama. Digital menghasilkan penjumlahan berat cup lump per lot pada tahap awal. RAM mengukur satu sistem yang berbeda: kendaraan dan muatan yang masih berada di dalamnya setelah pemindahan, penyusunan, penambahan waktu loading, dan perjalanan.",
    )
    base.add_callout(
        doc,
        "Konsekuensi metodologis",
        "Walaupun kedua alat bekerja sempurna, angka dapat berbeda karena perubahan fisik dan perubahan isi muatan terjadi di antara dua pengukuran. Karena itu, gap digital-RAM adalah selisih rantai proses atau neraca massa dispatch-to-RAM, bukan bukti langsung bahwa salah satu timbangan salah.",
        fill=base.PALE_GOLD,
        accent=base.GOLD,
    )
    base.add_body_paragraph(
        doc,
        "Arah digital yang selalu lebih tinggi sesuai dengan kemungkinan air/material keluar atau sebagian berat yang tercatat digital tidak lagi berada di truk saat RAM. Namun arah tersebut juga dapat muncul jika digital mempunyai bias ke atas atau net RAM mempunyai bias ke bawah. Data saat ini belum memisahkan komponen-komponen tersebut.",
    )

    base.add_section_title(doc, "04", "Apakah Selisihnya Harus Sebesar Itu?")
    add_magnitude_table(doc)
    base.add_body_paragraph(
        doc,
        "Penelitian Pusat Penelitian Karet terhadap slab dan lump melaporkan susut harian tertinggi 6,31% pada bahan berkadar karet kering rendah dan 2,41% pada lump mangkok berkadar karet kering lebih tinggi. Susut paling tinggi terjadi pada 1-2 hari pertama. Angka tersebut mendukung bahwa waktu dan kadar air dapat menghasilkan penurunan nyata, tetapi tidak boleh dipakai sebagai toleransi langsung untuk PT JP karena jenis, umur, kondisi, dan metode pengujiannya berbeda.",
    )
    base.add_body_paragraph(
        doc,
        "Gap 7,12%-21,17% secara fisik mungkin terjadi pada material yang mengandung banyak air, tetapi belum dapat langsung dianggap normal. Kasus 16,46% pada tanggal yang sama menunjukkan bahwa penjelasan 'menunggu sampai besok' tidak cukup. Kasus 21,17% memerlukan air/material sekitar 2.103 kg keluar atau kombinasi bias proses dan pengukuran dengan besaran setara.",
    )
    base.add_callout(
        doc,
        "Kesimpulan atas besaran gap",
        "Gapnya dapat diperkirakan akan ada, tetapi kewajaran besarannya belum terbukti. Nilai sampai 21,17% harus direkonsiliasi, bukan langsung diterima sebagai susut waktu dan bukan pula langsung dianggap kesalahan alat.",
        fill=base.PALE_RED,
        accent=base.RED,
    )

    base.add_section_title(doc, "05", "Pola yang Terlihat pada Empat Pasangan")
    add_pattern_tables(doc)
    add_relationship_table(doc)
    for text in [
        "Kelompok hari berikutnya mempunyai gap 16,20%, dibanding 11,49% pada hari yang sama. Kenaikannya 4,72 poin persentase, tetapi faktor kendaraan tidak terkontrol.",
        "Gembung mempunyai gap 18,84%, hampir dua kali Sebayur 9,06%. Pola sumber produksi lebih kuat daripada pola total muatan.",
        "Pada BK 8035 WS dengan kategori hari berikutnya, net digital 9.913 kg dan 9.933 kg hampir sama, tetapi gap berbeda 990 kg. Total muatan bukan penjelasan yang memadai.",
        "Korelasi dihitung hanya untuk deskripsi empat data. Dengan sampel sekecil ini, angkanya tidak dapat digunakan sebagai bukti statistik atau prediksi.",
    ]:
        base.add_bullet(doc, text)

    base.add_section_title(doc, "06", "Analisis Workbook Juli-Agustus")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(TIMELINE_PATH), width=Cm(16.4))
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    base.add_run(caption, "Gambar 3. Timeline Susut Gabungan sesuai label metode pada workbook.", size=8, color=base.MID_GREY, italic=True)
    base.add_stats_table(doc, source_rows)
    base.add_body_paragraph(
        doc,
        "Rata-rata sederhana Susut Gabungan adalah 38,93% untuk Juli-RAM, 33,74% untuk Agustus-RAM, dan 25,67% untuk Agustus-Digital. Median masing-masing 39,81%, 37,90%, dan 25,50%. Arah ini konsisten dengan net digital yang lebih tinggi sehingga susut terhitung lebih rendah.",
    )
    base.add_callout(
        doc,
        "Batas interpretasi Excel",
        "Perbedaan kelompok metode tidak membuktikan bahwa metode timbang menyebabkan perubahan persentase. Digital hanya memiliki tiga tanggal, RAM Agustus lima tanggal, volume DO tidak tersedia, dan pemilihan metode mengikuti kesiapan operasional sehingga bukan perbandingan acak atau berpasangan.",
        fill=base.PALE_GOLD,
        accent=base.GOLD,
    )

    landscape = doc.add_section(WD_SECTION.NEW_PAGE)
    base.configure_section(landscape, landscape=True)
    base.add_section_title(doc, "07", "Faktor yang Dapat Membentuk Selisih")
    add_factor_matrix(doc)
    base.add_callout(
        doc,
        "Cara membaca matriks",
        "Faktor-faktor di atas dapat terjadi bersamaan. Urutannya menunjukkan relevansi untuk diperiksa, bukan pembagian persentase penyebab dan bukan penetapan kesalahan pihak tertentu.",
        fill=base.PALE_GOLD,
        accent=base.GOLD,
    )

    portrait = doc.add_section(WD_SECTION.NEW_PAGE)
    base.configure_section(portrait, landscape=False)
    base.add_section_title(doc, "08", "Kesimpulan Berlapis")
    add_evidence_conclusions(doc)
    base.add_body_paragraph(
        doc,
        "Kesimpulan manajemen yang paling aman adalah bahwa PT JP sedang membandingkan dua basis berat yang berbeda: berat dispatch per lot pada saat digital dan berat truk setelah tahapan operasional pada saat RAM. Perbedaan memang diharapkan, tetapi toleransi yang dapat diterima belum dapat ditentukan dari empat pasangan ini.",
    )
    base.add_body_paragraph(
        doc,
        "Sebelum menentukan angka resmi, kedua nilai perlu dipertahankan beserta waktu dan sumbernya. Digital berfungsi sebagai bukti berat lot pada saat penimbangan awal, sedangkan RAM menjadi bukti berat muatan pada titik berikutnya. Gap diperlakukan sebagai objek rekonsiliasi sampai komponen material, proses, tare, kendaraan, dan alat dapat diukur.",
    )

    base.add_section_title(doc, "09", "Data yang Dibutuhkan untuk Menjawab Besaran Gap")
    add_required_data_table(doc)
    base.add_callout(
        doc,
        "Uji pembeda yang disarankan",
        "Lakukan pengujian terkontrol pada kendaraan dan produksi yang sama: timbang lot digital, catat dan tampung air/residu yang tertinggal, rekonsiliasi seluruh lot ke truk, gunakan tare aktual, lalu timbang RAM dengan timestamp lengkap. Pengulangan silang BD/BK dan Gembung/Sebayur akan memisahkan efek waktu, material, kendaraan, dan alat.",
    )

    doc.add_page_break()
    base.add_section_title(doc, "A", "Lampiran Data Workbook")
    base.add_body_paragraph(
        doc,
        "Seluruh nilai berikut disalin dari sheet Ekstrak Susut. Tanda '-' menggantikan nilai nol pada produksi yang tidak memiliki angka susut pada tanggal tersebut.",
    )
    base.add_source_appendix(doc, source_rows)

    base.add_section_title(doc, "B", "Definisi Perhitungan")
    definitions = [
        ("Gap kilogram", "Net digital - Net RAM"),
        ("Gap persen", "(Net digital - Net RAM) / Net digital x 100%"),
        ("Gap kelompok", "Jumlah gap / jumlah net digital dalam kelompok"),
        ("Rata-rata workbook", "Rata-rata sederhana antar tanggal; tidak berbobot tonase"),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for index, header in enumerate(("Istilah", "Definisi")):
        base.write_cell(table.cell(0, index), header, bold=True, size=8.5, color=base.WHITE, align=WD_ALIGN_PARAGRAPH.CENTER, fill=base.GREEN)
    for row_index, values in enumerate(definitions, start=1):
        fill = base.WHITE if row_index % 2 else base.VERY_LIGHT_GREY
        cells = table.add_row().cells
        base.write_cell(cells[0], values[0], bold=True, size=8.5, color=base.GREEN_2, fill=fill)
        base.write_cell(cells[1], values[1], size=8.5, fill=fill)
    base.set_table_borders(table, base.LIGHT_GREY, 5)

    base.add_section_title(doc, "C", "Referensi")
    references = [
        "Rachmawan, A., & Wijaya, A. (2018). Pengaruh Kadar Karet Kering Lateks pada Susut Bobot Slab dan Lump. Warta Perkaretan, 37(1), 51-60. https://doi.org/10.22302/ppk.wp.v37i1.556",
        "Nair, D. B., Jacob, J., & Nair, N. R. (2012). A simple method for rapid determination of residual water content in rubber cup lumps. Journal of Plantation Crops, 40(1), 35-39. https://updatepublishing.com/journal/index.php/JPC/article/view/7575",
        "Australian National Measurement Institute. Weighbridge Operators Manual. https://www.industry.gov.au/sites/default/files/2019-03/weighbridgeoperatorsmanual.pdf",
        "International Organization of Legal Metrology. (2013). Application of truck scales in the petroleum industry. OIML Bulletin, LIV(3). https://www.oiml.org/en/publications/bulletin/pdf/oiml_bulletin_july_2013.pdf",
        "PT Julang Plantations. (2026). Screenshot perbandingan empat DO dan workbook perbandingan susut RAM dan digital.",
    ]
    for ref in references:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.first_line_indent = Cm(-0.5)
        base.add_run(p, ref, size=8.8, color=base.CHARCOAL)

    doc.core_properties.title = "Laporan Analisis Berbasis Data Digital Per Lot vs RAM"
    doc.core_properties.subject = "Analisis empat pasangan timbang dan workbook susut Juli-Agustus 2026"
    doc.core_properties.author = "PT Julang Plantations"
    doc.core_properties.keywords = "timbangan digital, timbangan RAM, cup lump, DO, selisih berat, susut"
    doc.core_properties.comments = "Dokumen baru yang berfokus pada data; tidak menetapkan penyebab tunggal."
    doc.core_properties.created = datetime(2026, 8, 26)
    doc.save(OUTPUT_DOCX)
    return OUTPUT_DOCX


if __name__ == "__main__":
    print(create_document())
