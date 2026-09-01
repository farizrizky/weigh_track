from __future__ import annotations

import math
from datetime import datetime
from pathlib import Path

import openpyxl
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE_XLSX = Path(r"C:\Users\USER\Downloads\perbandingan susut ram dan digital.xlsx")
OUTPUT_DIR = ROOT / "reports"
ASSET_DIR = OUTPUT_DIR / "assets"
OUTPUT_DOCX = OUTPUT_DIR / "Laporan_Analisis_Sementara_Timbangan_Digital_vs_RAM.docx"
CHART_PATH = ASSET_DIR / "perbandingan_net_digital_ram.png"

GREEN = "1F4E3D"
GREEN_2 = "3E755E"
LIGHT_GREEN = "DCE8E0"
PALE_GREEN = "EEF4F0"
GOLD = "C79A3B"
PALE_GOLD = "F7F0DE"
CHARCOAL = "252B29"
MID_GREY = "68716D"
LIGHT_GREY = "E7EBE9"
VERY_LIGHT_GREY = "F5F7F6"
WHITE = "FFFFFF"
RED = "A13A32"
PALE_RED = "F7E8E6"


PAIRED_ROWS = [
    {
        "no": 1,
        "digital_date": "11 Agu 2026",
        "ram_date": "11 Agu 2026",
        "origin": "Gembung",
        "truck": "BD 8450 K",
        "digital": 9732,
        "ram": 8130,
        "timing": "Hari yang sama",
    },
    {
        "no": 2,
        "digital_date": "13 Agu 2026",
        "ram_date": "14 Agu 2026",
        "origin": "Sebayur",
        "truck": "BK 8035 WS",
        "digital": 9913,
        "ram": 8800,
        "timing": "Hari berikutnya",
    },
    {
        "no": 3,
        "digital_date": "24 Agu 2026",
        "ram_date": "24 Agu 2026",
        "origin": "Sebayur",
        "truck": "BD 8450 K",
        "digital": 11090,
        "ram": 10300,
        "timing": "Hari yang sama",
    },
    {
        "no": 4,
        "digital_date": "24 Agu 2026",
        "ram_date": "25 Agu 2026",
        "origin": "Gembung",
        "truck": "BK 8035 WS",
        "digital": 9933,
        "ram": 7830,
        "timing": "Hari berikutnya",
    },
]

for row in PAIRED_ROWS:
    row["gap"] = row["digital"] - row["ram"]
    row["gap_pct"] = row["gap"] / row["digital"]


def format_kg(value: float | int) -> str:
    return f"{int(round(value)):,}".replace(",", ".")


def format_pct(value: float, decimals: int = 2) -> str:
    return f"{value * 100:.{decimals}f}%".replace(".", ",")


def rgb(hex_color: str) -> RGBColor:
    return RGBColor.from_string(hex_color)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, **kwargs) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)

    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        edge_data = kwargs.get(edge)
        if not edge_data:
            continue
        tag = "w:" + edge
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        for key in ("val", "sz", "space", "color"):
            if key in edge_data:
                element.set(qn("w:" + key), str(edge_data[key]))


def set_cell_margins(cell, top=80, start=90, bottom=80, end=90) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn("w:" + margin))
        if node is None:
            node = OxmlElement("w:" + margin)
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_cell_width(cell, width_cm: float) -> None:
    cell.width = Cm(width_cm)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(Cm(width_cm).twips))
    tc_w.set(qn("w:type"), "dxa")


def set_table_fixed_layout(table) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")


def set_table_borders(table, color=LIGHT_GREY, size=6) -> None:
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(
                cell,
                top={"val": "single", "sz": size, "color": color},
                bottom={"val": "single", "sz": size, "color": color},
                left={"val": "single", "sz": size, "color": color},
                right={"val": "single", "sz": size, "color": color},
            )


def add_run(paragraph, text: str, *, bold=False, size=None, color=None, italic=False):
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = rgb(color)
    run.font.name = "Aptos"
    return run


def clear_cell(cell) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)


def write_cell(
    cell,
    text: str,
    *,
    bold=False,
    size=8.5,
    color=CHARCOAL,
    align=WD_ALIGN_PARAGRAPH.LEFT,
    fill=None,
    valign=WD_CELL_VERTICAL_ALIGNMENT.CENTER,
) -> None:
    clear_cell(cell)
    cell.vertical_alignment = valign
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    add_run(p, str(text), bold=bold, size=size, color=color)
    set_cell_margins(cell)
    if fill:
        set_cell_shading(cell, fill)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_run(paragraph, "Halaman ", size=8, color=MID_GREY)
    run = paragraph.add_run()
    fld_char_1 = OxmlElement("w:fldChar")
    fld_char_1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char_2 = OxmlElement("w:fldChar")
    fld_char_2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_1)
    run._r.append(instr_text)
    run._r.append(fld_char_2)


def configure_section(section, landscape=False) -> None:
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.7)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.header_distance = Cm(0.7)
    section.footer_distance = Cm(0.7)
    if landscape:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Cm(29.7)
        section.page_height = Cm(21.0)
    else:
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)


def configure_header_footer(section) -> None:
    header = section.header
    table = header.add_table(rows=1, cols=2, width=Cm(17.0))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_fixed_layout(table)
    set_cell_width(table.cell(0, 0), 8.5)
    set_cell_width(table.cell(0, 1), 8.5)
    write_cell(table.cell(0, 0), "PT JULANG PLANTATIONS", bold=True, size=8.5, color=GREEN)
    write_cell(
        table.cell(0, 1),
        "ANALISIS OPERASIONAL | DISTRIBUSI KARET",
        bold=True,
        size=7.5,
        color=MID_GREY,
        align=WD_ALIGN_PARAGRAPH.RIGHT,
    )
    for cell in table.rows[0].cells:
        set_cell_border(cell, bottom={"val": "single", "sz": 14, "color": GREEN})

    footer = section.footer
    table = footer.add_table(rows=1, cols=2, width=Cm(17.0))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_fixed_layout(table)
    write_cell(
        table.cell(0, 0),
        "Laporan Analisis Sementara - 26 Agustus 2026",
        size=7.5,
        color=MID_GREY,
    )
    clear_cell(table.cell(0, 1))
    add_page_number(table.cell(0, 1).paragraphs[0])
    for cell in table.rows[0].cells:
        set_cell_border(cell, top={"val": "single", "sz": 8, "color": LIGHT_GREY})


def setup_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10)
    normal.font.color.rgb = rgb(CHARCOAL)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.13

    for style_name, size, color in (
        ("Title", 25, GREEN),
        ("Subtitle", 12, MID_GREY),
        ("Heading 1", 17, GREEN),
        ("Heading 2", 12.5, GREEN_2),
        ("Heading 3", 10.5, CHARCOAL),
    ):
        style = doc.styles[style_name]
        style.font.name = "Aptos Display" if style_name != "Normal" else "Aptos"
        style.font.size = Pt(size)
        style.font.color.rgb = rgb(color)
        style.font.bold = style_name != "Subtitle"
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.space_before = Pt(10 if style_name != "Title" else 0)
        style.paragraph_format.space_after = Pt(5)

    doc.styles["Heading 1"].paragraph_format.page_break_before = False


def add_section_title(doc: Document, number: str, title: str, subtitle: str | None = None) -> None:
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 1"]
    p.paragraph_format.space_before = Pt(14)
    add_run(p, number, bold=True, size=11, color=GOLD)
    add_run(p, "  " + title, bold=True, size=17, color=GREEN)
    if subtitle:
        p2 = doc.add_paragraph()
        p2.paragraph_format.space_after = Pt(8)
        add_run(p2, subtitle, size=9, color=MID_GREY, italic=True)


def add_body_paragraph(doc: Document, text: str, *, bold_lead: str | None = None, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.13
    if bold_lead and text.startswith(bold_lead):
        add_run(p, bold_lead, bold=True, size=10, color=CHARCOAL)
        add_run(p, text[len(bold_lead) :], size=10, color=CHARCOAL, italic=italic)
    else:
        add_run(p, text, size=10, color=CHARCOAL, italic=italic)
    return p


def add_bullet(doc: Document, text: str, *, level=0, color=CHARCOAL) -> None:
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.55 + level * 0.45)
    p.paragraph_format.first_line_indent = Cm(-0.25)
    add_run(p, text, size=9.5, color=color)


def add_number(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(0.65)
    p.paragraph_format.first_line_indent = Cm(-0.35)
    add_run(p, text, size=9.5, color=CHARCOAL)


def add_callout(doc: Document, title: str, body: str, *, fill=PALE_GREEN, accent=GREEN) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_fixed_layout(table)
    set_cell_width(table.cell(0, 0), 0.22)
    set_cell_width(table.cell(0, 1), 16.5)
    set_cell_shading(table.cell(0, 0), accent)
    set_cell_border(table.cell(0, 0), top={"val": "nil"}, bottom={"val": "nil"})
    cell = table.cell(0, 1)
    set_cell_shading(cell, fill)
    clear_cell(cell)
    p = cell.paragraphs[0]
    add_run(p, title, bold=True, size=9.5, color=accent)
    p.paragraph_format.space_after = Pt(2)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.08
    add_run(p2, body, size=9.2, color=CHARCOAL)
    set_cell_margins(cell, top=130, start=170, bottom=130, end=170)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_kpi_cards(doc: Document) -> None:
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_fixed_layout(table)
    cards = [
        ("4 DO", "Perbandingan langsung"),
        ("5.608 kg", "Total selisih net"),
        ("7,12%-21,17%", "Rentang penurunan"),
    ]
    for index, (value, label) in enumerate(cards):
        cell = table.cell(0, index)
        set_cell_shading(cell, VERY_LIGHT_GREY if index != 1 else PALE_GOLD)
        clear_cell(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, value, bold=True, size=15, color=GREEN if index != 1 else GOLD)
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(0)
        add_run(p2, label, size=8.5, color=MID_GREY)
        set_cell_margins(cell, top=180, start=100, bottom=180, end=100)
    set_table_borders(table, WHITE, 4)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_process_flow(doc: Document) -> None:
    add_body_paragraph(doc, "Proses sebelum timbangan digital", bold_lead="Proses sebelum timbangan digital")
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_fixed_layout(table)
    labels = ["1. Muat langsung", "2. Perjalanan ke RAM", "3. Timbang gross truk", "4. Net = gross - tare"]
    for i, label in enumerate(labels):
        write_cell(
            table.cell(0, i),
            label,
            bold=True,
            size=8.5,
            color=WHITE if i in (0, 2) else CHARCOAL,
            align=WD_ALIGN_PARAGRAPH.CENTER,
            fill=GREEN_2 if i in (0, 2) else LIGHT_GREEN,
        )
    set_table_borders(table, WHITE, 8)

    add_body_paragraph(doc, "Proses setelah timbangan digital", bold_lead="Proses setelah timbangan digital")
    table = doc.add_table(rows=1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_fixed_layout(table)
    labels = [
        "1. Timbang lot",
        "2. Akumulasi net",
        "3. Loading +0,5-1 jam",
        "4. Perjalanan",
        "5. Timbang RAM",
    ]
    for i, label in enumerate(labels):
        write_cell(
            table.cell(0, i),
            label,
            bold=True,
            size=8.2,
            color=WHITE if i in (0, 2, 4) else CHARCOAL,
            align=WD_ALIGN_PARAGRAPH.CENTER,
            fill=GREEN_2 if i in (0, 2, 4) else LIGHT_GREEN,
        )
    set_table_borders(table, WHITE, 8)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def make_chart() -> None:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    width, height = 1600, 820
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)

    def font(size: int, bold=False):
        candidates = [
            Path(r"C:\Windows\Fonts\arialbd.ttf" if bold else r"C:\Windows\Fonts\arial.ttf"),
            Path(r"C:\Windows\Fonts\calibrib.ttf" if bold else r"C:\Windows\Fonts\calibri.ttf"),
        ]
        for candidate in candidates:
            if candidate.exists():
                return ImageFont.truetype(str(candidate), size=size)
        return ImageFont.load_default()

    title_font = font(38, bold=True)
    subtitle_font = font(22)
    axis_font = font(21)
    label_font = font(22, bold=True)
    small_font = font(18)

    draw.text((90, 45), "Perbandingan Net Digital dan RAM", font=title_font, fill="#1F4E3D")
    draw.text(
        (90, 96),
        "Empat DO yang memiliki pasangan pengukuran | satuan kilogram",
        font=subtitle_font,
        fill="#68716D",
    )

    chart_left, chart_top, chart_right, chart_bottom = 120, 165, 1530, 670
    max_value = 12000
    tick_step = 2000
    for tick in range(0, max_value + 1, tick_step):
        y = chart_bottom - (tick / max_value) * (chart_bottom - chart_top)
        draw.line((chart_left, y, chart_right, y), fill="#E7EBE9", width=2)
        label = format_kg(tick)
        bbox = draw.textbbox((0, 0), label, font=small_font)
        draw.text((chart_left - 18 - (bbox[2] - bbox[0]), y - 10), label, font=small_font, fill="#68716D")

    group_width = (chart_right - chart_left) / len(PAIRED_ROWS)
    bar_width = 100
    gap_between = 22
    for index, row in enumerate(PAIRED_ROWS):
        center = chart_left + group_width * (index + 0.5)
        values = [(row["digital"], "#3E755E"), (row["ram"], "#9CA7A2")]
        x_positions = [center - bar_width - gap_between / 2, center + gap_between / 2]
        for (value, color), x in zip(values, x_positions):
            y = chart_bottom - (value / max_value) * (chart_bottom - chart_top)
            draw.rounded_rectangle((x, y, x + bar_width, chart_bottom), radius=6, fill=color)
            value_text = format_kg(value)
            bbox = draw.textbbox((0, 0), value_text, font=label_font)
            draw.text(
                (x + (bar_width - (bbox[2] - bbox[0])) / 2, y - 34),
                value_text,
                font=label_font,
                fill="#252B29",
            )
        x_label = f"DO {row['no']} | {row['origin']}"
        bbox = draw.textbbox((0, 0), x_label, font=axis_font)
        draw.text(
            (center - (bbox[2] - bbox[0]) / 2, chart_bottom + 24),
            x_label,
            font=axis_font,
            fill="#252B29",
        )
        gap_label = f"Selisih {format_pct(row['gap_pct'])}"
        bbox = draw.textbbox((0, 0), gap_label, font=small_font)
        draw.text(
            (center - (bbox[2] - bbox[0]) / 2, chart_bottom + 58),
            gap_label,
            font=small_font,
            fill="#A13A32",
        )

    legend_y = 755
    draw.rounded_rectangle((610, legend_y, 645, legend_y + 24), radius=4, fill="#3E755E")
    draw.text((657, legend_y - 3), "Digital", font=axis_font, fill="#252B29")
    draw.rounded_rectangle((790, legend_y, 825, legend_y + 24), radius=4, fill="#9CA7A2")
    draw.text((837, legend_y - 3), "RAM", font=axis_font, fill="#252B29")

    image.save(CHART_PATH, quality=95)


def load_source_rows():
    workbook = openpyxl.load_workbook(SOURCE_XLSX, data_only=True)
    sheet = workbook["Ekstrak Susut"]
    rows = []
    for values in sheet.iter_rows(min_row=2, values_only=True):
        date_value, sebayur, gembung, combined, scale = values
        rows.append(
            {
                "date": date_value,
                "sebayur": float(sebayur or 0),
                "gembung": float(gembung or 0),
                "combined": float(combined or 0),
                "scale": str(scale),
            }
        )
    return rows


def basic_stats(values):
    ordered = sorted(values)
    count = len(ordered)
    mean = sum(ordered) / count
    if count % 2:
        median = ordered[count // 2]
    else:
        median = (ordered[count // 2 - 1] + ordered[count // 2]) / 2
    return {
        "n": count,
        "mean": mean,
        "median": median,
        "min": min(ordered),
        "max": max(ordered),
    }


def add_paired_table(doc: Document) -> None:
    headers = [
        "No.",
        "Tanggal Digital / RAM",
        "Produksi",
        "Nopol",
        "Net Digital",
        "Net RAM",
        "Selisih",
        "% dari Digital",
    ]
    widths = [0.7, 2.8, 1.6, 1.9, 1.8, 1.7, 1.6, 1.8]
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_fixed_layout(table)
    for index, header in enumerate(headers):
        write_cell(
            table.cell(0, index),
            header,
            bold=True,
            size=7.5,
            color=WHITE,
            align=WD_ALIGN_PARAGRAPH.CENTER,
            fill=GREEN,
        )
        set_cell_width(table.cell(0, index), widths[index])
    set_repeat_table_header(table.rows[0])

    for row_index, row in enumerate(PAIRED_ROWS, start=1):
        values = [
            row["no"],
            f"{row['digital_date']}\n{row['ram_date']}",
            row["origin"],
            row["truck"],
            f"{format_kg(row['digital'])} kg",
            f"{format_kg(row['ram'])} kg",
            f"{format_kg(row['gap'])} kg",
            format_pct(row["gap_pct"]),
        ]
        cells = table.add_row().cells
        fill = WHITE if row_index % 2 else VERY_LIGHT_GREY
        for index, value in enumerate(values):
            align = WD_ALIGN_PARAGRAPH.RIGHT if index >= 4 else WD_ALIGN_PARAGRAPH.CENTER
            write_cell(cells[index], value, size=7.8, align=align, fill=fill)
            set_cell_width(cells[index], widths[index])
        write_cell(cells[6], values[6], bold=True, size=7.8, color=RED, align=WD_ALIGN_PARAGRAPH.RIGHT, fill=fill)
        write_cell(cells[7], values[7], bold=True, size=7.8, color=RED, align=WD_ALIGN_PARAGRAPH.RIGHT, fill=fill)
        prevent_row_split(table.rows[-1])

    total_digital = sum(row["digital"] for row in PAIRED_ROWS)
    total_ram = sum(row["ram"] for row in PAIRED_ROWS)
    total_gap = total_digital - total_ram
    cells = table.add_row().cells
    merged = cells[0].merge(cells[3])
    write_cell(merged, "TOTAL EMPAT DO", bold=True, size=8, color=WHITE, align=WD_ALIGN_PARAGRAPH.RIGHT, fill=GREEN_2)
    write_cell(cells[4], f"{format_kg(total_digital)} kg", bold=True, size=8, color=WHITE, align=WD_ALIGN_PARAGRAPH.RIGHT, fill=GREEN_2)
    write_cell(cells[5], f"{format_kg(total_ram)} kg", bold=True, size=8, color=WHITE, align=WD_ALIGN_PARAGRAPH.RIGHT, fill=GREEN_2)
    write_cell(cells[6], f"{format_kg(total_gap)} kg", bold=True, size=8, color=WHITE, align=WD_ALIGN_PARAGRAPH.RIGHT, fill=RED)
    write_cell(cells[7], format_pct(total_gap / total_digital), bold=True, size=8, color=WHITE, align=WD_ALIGN_PARAGRAPH.RIGHT, fill=RED)
    set_table_borders(table, WHITE, 5)


def add_stats_table(doc: Document, source_rows) -> None:
    groups = [
        ("Juli 2026 - RAM", [r["combined"] for r in source_rows if r["date"].month == 7 and r["scale"].lower() == "ram"]),
        ("Agustus 2026 - RAM", [r["combined"] for r in source_rows if r["date"].month == 8 and r["scale"].lower() == "ram"]),
        (
            "Agustus 2026 - Digital",
            [r["combined"] for r in source_rows if r["date"].month == 8 and r["scale"].lower() == "digital"],
        ),
    ]
    headers = ["Kelompok data", "Jumlah tanggal", "Rata-rata", "Median", "Minimum", "Maksimum"]
    widths = [4.3, 2.1, 2.2, 2.2, 2.2, 2.2]
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_fixed_layout(table)
    for index, header in enumerate(headers):
        write_cell(table.cell(0, index), header, bold=True, size=8, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER, fill=GREEN)
        set_cell_width(table.cell(0, index), widths[index])
    set_repeat_table_header(table.rows[0])
    for row_index, (label, values) in enumerate(groups, start=1):
        stats = basic_stats(values)
        row_values = [
            label,
            stats["n"],
            format_pct(stats["mean"]),
            format_pct(stats["median"]),
            format_pct(stats["min"]),
            format_pct(stats["max"]),
        ]
        cells = table.add_row().cells
        fill = WHITE if row_index % 2 else VERY_LIGHT_GREY
        for index, value in enumerate(row_values):
            align = WD_ALIGN_PARAGRAPH.LEFT if index == 0 else WD_ALIGN_PARAGRAPH.CENTER
            write_cell(cells[index], value, size=8.2, align=align, fill=fill)
            set_cell_width(cells[index], widths[index])
        prevent_row_split(table.rows[-1])
    set_table_borders(table, LIGHT_GREY, 5)


def add_evidence_matrix(doc: Document) -> None:
    headers = ["Prioritas", "Hipotesis", "Mekanisme", "Indikasi saat ini", "Cara membuktikan"]
    widths = [1.4, 4.2, 6.2, 6.4, 6.5]
    rows = [
        (
            "1",
            "Air keluar akibat tekanan tumpukan, waktu, dan getaran",
            "Lot ditimbang saat belum tertekan; tumpukan di truk mendorong air/serum keluar sebelum RAM.",
            "Digital selalu lebih tinggi; kasus hari yang sama tetap turun 7,12%-16,46%. Gembung memiliki gap lebih besar.",
            "Uji timbang ulang 0/30/60/90 menit dan tampung cairan yang keluar.",
        ),
        (
            "2",
            "Lot tertimbang tidak sama dengan lot termuat",
            "Lot tertinggal, berpindah truk, batal tetapi terjumlah, duplikat, atau salah DO.",
            "Digital merupakan agregasi banyak kejadian; kesalahan per lot dapat terakumulasi.",
            "Rekonsiliasi ID lot, jumlah fisik, status batal, waktu timbang, dan checklist loading.",
        ),
        (
            "3",
            "Basis gross digital tidak setara dengan gross RAM",
            "Gross digital mungkin dihitung dari jumlah net lot + tare, bukan pembacaan truk secara langsung.",
            "Alur yang dijelaskan adalah timbang lot, sedangkan RAM menimbang truk final.",
            "Konfirmasi definisi kolom dan telusuri sumber setiap angka pada empat DO.",
        ),
        (
            "4",
            "Tare kendaraan tidak aktual",
            "Tare historis dipakai walaupun bahan bakar, pengemudi, terpal, alat, lumpur, atau air berubah.",
            "Tare identik pada setiap pasangan sehingga gap tabel berasal dari gross, tetapi net RAM tetap dapat bias.",
            "Timbang kosong kendaraan pada perjalanan yang sama dan catat konfigurasi truk.",
        ),
        (
            "5",
            "Bias alat atau prosedur timbang",
            "Zero drift, tare wadah, indikator belum stabil, instalasi, posisi truk, atau kalibrasi.",
            "Rentang gap 7,12%-21,17% tidak menyerupai satu offset tetap, tetapi bias sistematis belum tersingkirkan.",
            "Uji beban standar, cek sertifikat, zero, repeatability, dan prosedur operator.",
        ),
        (
            "6",
            "Kehilangan fisik setelah loading",
            "Air, potongan karet, atau material lain keluar selama perjalanan; muatan dapat berubah.",
            "Belum tersedia dokumentasi segel, kondisi bak, atau volume runoff.",
            "Gunakan segel, foto, checklist, dan inspeksi bak sebelum/sesudah perjalanan.",
        ),
    ]

    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_fixed_layout(table)
    for index, header in enumerate(headers):
        write_cell(table.cell(0, index), header, bold=True, size=8.0, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER, fill=GREEN)
        set_cell_width(table.cell(0, index), widths[index])
    set_repeat_table_header(table.rows[0])
    for row_index, values in enumerate(rows, start=1):
        cells = table.add_row().cells
        fill = WHITE if row_index % 2 else VERY_LIGHT_GREY
        for index, value in enumerate(values):
            align = WD_ALIGN_PARAGRAPH.CENTER if index == 0 else WD_ALIGN_PARAGRAPH.LEFT
            write_cell(cells[index], value, size=8.0, align=align, fill=fill, valign=WD_CELL_VERTICAL_ALIGNMENT.TOP)
            set_cell_width(cells[index], widths[index])
        write_cell(cells[0], values[0], bold=True, size=9, color=GREEN, align=WD_ALIGN_PARAGRAPH.CENTER, fill=fill)
        prevent_row_split(table.rows[-1])
    set_table_borders(table, LIGHT_GREY, 5)


def add_compression_volume_table(doc: Document) -> None:
    headers = ["DO", "Produksi", "Selisih", "Ekuivalen cairan*"]
    widths = [2.0, 3.0, 3.0, 4.2]
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_fixed_layout(table)
    for index, header in enumerate(headers):
        write_cell(table.cell(0, index), header, bold=True, size=8, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER, fill=GREEN)
        set_cell_width(table.cell(0, index), widths[index])
    for row_index, row in enumerate(PAIRED_ROWS, start=1):
        cells = table.add_row().cells
        fill = WHITE if row_index % 2 else VERY_LIGHT_GREY
        values = [row["no"], row["origin"], f"{format_kg(row['gap'])} kg", f"sekitar {format_kg(row['gap'])} liter"]
        for index, value in enumerate(values):
            align = WD_ALIGN_PARAGRAPH.CENTER if index < 2 else WD_ALIGN_PARAGRAPH.RIGHT
            write_cell(cells[index], value, size=8.2, align=align, fill=fill)
            set_cell_width(cells[index], widths[index])
    set_table_borders(table, LIGHT_GREY, 5)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    add_run(
        p,
        "*Pendekatan menggunakan massa jenis cairan mendekati air, sekitar 1 kg/liter. Ini indikator kewajaran, bukan hasil pengukuran runoff.",
        size=8,
        color=MID_GREY,
        italic=True,
    )


def add_validation_plan(doc: Document) -> None:
    phases = [
        (
            "TAHAP 1",
            "Audit empat DO yang sudah berpasangan",
            "Selesaikan 1-2 hari",
            [
                "Konfirmasi apakah gross digital merupakan pembacaan alat atau rekonstruksi jumlah net lot + tare.",
                "Rekonsiliasi seluruh ID lot, status batal, duplikat, perpindahan DO, dan checklist fisik loading.",
                "Periksa tiket RAM asli, sumber tare, waktu timbang, operator, dan perubahan kendaraan.",
                "Audit kemunculan gross digital 14.633 kg pada dua DO berbeda untuk BK 8035 WS.",
            ],
        ),
        (
            "TAHAP 2",
            "Uji tekanan tumpukan skala terbatas",
            "Minimal 3 batch representatif",
            [
                "Ambil batch 500-1.000 kg dari Sebayur dan Gembung; timbang pada digital sebelum ditumpuk.",
                "Berikan tekanan/tumpukan yang menyerupai bagian bawah bak truk.",
                "Tampung cairan dan timbang ulang pada menit ke-30, 60, dan 90.",
                "Bandingkan kehilangan berat dengan cairan tertampung; catat hujan, usia karet, ukuran lump, dan kondisi basah.",
            ],
        ),
        (
            "TAHAP 3",
            "Uji perjalanan terkontrol",
            "5-10 DO dengan desain silang",
            [
                "Gunakan kedua kendaraan secara bergantian untuk Sebayur dan Gembung serta variasikan hari yang sama/hari berikutnya.",
                "Catat waktu timbang lot pertama/terakhir, selesai loading, berangkat, dan timbang RAM.",
                "Gunakan tare kosong aktual, segel kendaraan, foto bak, serta pencatatan cairan yang keluar.",
                "Lakukan pemeriksaan zero dan beban standar sebelum operasi; minta bukti kalibrasi RAM.",
            ],
        ),
    ]

    for label, title, target, bullets in phases:
        table = doc.add_table(rows=1, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_table_fixed_layout(table)
        set_cell_width(table.cell(0, 0), 2.2)
        set_cell_width(table.cell(0, 1), 14.4)
        write_cell(table.cell(0, 0), label, bold=True, size=9, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER, fill=GREEN)
        cell = table.cell(0, 1)
        set_cell_shading(cell, VERY_LIGHT_GREY)
        clear_cell(cell)
        p = cell.paragraphs[0]
        add_run(p, title, bold=True, size=10, color=GREEN)
        add_run(p, "  |  " + target, size=8.5, color=MID_GREY, italic=True)
        p.paragraph_format.space_after = Pt(4)
        for bullet in bullets:
            bp = cell.add_paragraph()
            bp.style = doc.styles["List Bullet"]
            bp.paragraph_format.space_after = Pt(2)
            bp.paragraph_format.left_indent = Cm(0.45)
            bp.paragraph_format.first_line_indent = Cm(-0.25)
            add_run(bp, bullet, size=8.8, color=CHARCOAL)
        set_cell_margins(cell, top=130, start=170, bottom=120, end=170)
        set_table_borders(table, WHITE, 5)
        doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_decision_rules(doc: Document) -> None:
    rows = [
        ("Berat batch turun dan cairan tertampung mendekati penurunan", "Tekanan tumpukan dan drainase merupakan penyebab dominan."),
        ("Jumlah digital melebihi lot yang terkonfirmasi naik truk", "Masalah utama berada pada rekonsiliasi loading/data."),
        ("Beban standar menunjukkan deviasi berulang", "Kalibrasi atau instalasi alat perlu dikoreksi."),
        ("Tare aktual berbeda material dari tare referensi", "Net RAM perlu dihitung ulang menggunakan tare aktual."),
        ("Cairan sedikit tetapi gap tetap besar", "Efek tekanan tidak cukup; lanjutkan audit lot, tare, dan kedua alat."),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_fixed_layout(table)
    for index, header in enumerate(("Hasil pengujian", "Interpretasi")):
        write_cell(table.cell(0, index), header, bold=True, size=8.5, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER, fill=GREEN)
    for row_index, values in enumerate(rows, start=1):
        cells = table.add_row().cells
        fill = WHITE if row_index % 2 else VERY_LIGHT_GREY
        write_cell(cells[0], values[0], size=8.5, fill=fill)
        write_cell(cells[1], values[1], bold=True, size=8.5, color=GREEN_2, fill=fill)
        prevent_row_split(table.rows[-1])
    set_table_borders(table, LIGHT_GREY, 5)


def add_source_appendix(doc: Document, source_rows) -> None:
    headers = ["Tanggal DO", "Susut Sebayur", "Susut Gembung", "Susut Gabungan", "Timbangan"]
    widths = [3.2, 3.0, 3.0, 3.2, 2.6]
    table = doc.add_table(rows=1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_fixed_layout(table)
    for index, header in enumerate(headers):
        write_cell(table.cell(0, index), header, bold=True, size=8, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER, fill=GREEN)
        set_cell_width(table.cell(0, index), widths[index])
    set_repeat_table_header(table.rows[0])
    for row_index, row in enumerate(source_rows, start=1):
        fill = WHITE if row_index % 2 else VERY_LIGHT_GREY
        cells = table.add_row().cells
        values = [
            row["date"].strftime("%d %b %Y").replace("Jul", "Jul").replace("Aug", "Agu"),
            "-" if row["sebayur"] == 0 else format_pct(row["sebayur"], 1),
            "-" if row["gembung"] == 0 else format_pct(row["gembung"], 1),
            format_pct(row["combined"], 1),
            row["scale"],
        ]
        for index, value in enumerate(values):
            align = WD_ALIGN_PARAGRAPH.CENTER
            write_cell(cells[index], value, size=8.2, align=align, fill=fill)
            set_cell_width(cells[index], widths[index])
        prevent_row_split(table.rows[-1])
    set_table_borders(table, LIGHT_GREY, 5)


def add_recommended_data_fields(doc: Document) -> None:
    headers = ["Kelompok", "Kolom minimum", "Tujuan"]
    rows = [
        ("Identitas", "DO, tanggal, produksi, nopol, operator, pengemudi", "Memisahkan pengaruh sumber, orang, dan kendaraan."),
        ("Digital", "Waktu tiap lot, ID lot, gross, tare wadah, net, status stabil", "Membuktikan jumlah transaksi dan kualitas pembacaan."),
        ("Loading", "Checklist lot termuat, posisi tumpukan, selesai loading, segel", "Memastikan lot tertimbang benar-benar terkirim."),
        ("RAM", "Waktu timbang, gross, tare aktual, net, tiket, operator", "Memastikan dasar net RAM dapat direproduksi."),
        ("Kondisi", "Hujan, kondisi basah, waktu tunggu, runoff, kondisi bak", "Mengukur pengaruh air, tekanan tumpukan, dan perjalanan."),
        ("Rekonsiliasi", "Gap kg, gap %, klasifikasi penyebab, bukti pendukung", "Mengubah selisih menjadi temuan yang dapat ditindaklanjuti."),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_fixed_layout(table)
    widths = [2.6, 7.4, 6.0]
    for index, header in enumerate(headers):
        write_cell(table.cell(0, index), header, bold=True, size=8, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER, fill=GREEN)
        set_cell_width(table.cell(0, index), widths[index])
    for row_index, values in enumerate(rows, start=1):
        cells = table.add_row().cells
        fill = WHITE if row_index % 2 else VERY_LIGHT_GREY
        for index, value in enumerate(values):
            write_cell(cells[index], value, size=8.2, fill=fill, valign=WD_CELL_VERTICAL_ALIGNMENT.TOP)
            set_cell_width(cells[index], widths[index])
        write_cell(cells[0], values[0], bold=True, size=8.2, color=GREEN_2, fill=fill, valign=WD_CELL_VERTICAL_ALIGNMENT.TOP)
        prevent_row_split(table.rows[-1])
    set_table_borders(table, LIGHT_GREY, 5)


def create_document() -> Path:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    source_rows = load_source_rows()
    make_chart()

    doc = Document()
    setup_styles(doc)
    configure_section(doc.sections[0])
    configure_header_footer(doc.sections[0])

    # Cover
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(32)
    add_run(p, "PT JULANG PLANTATIONS", bold=True, size=11, color=GREEN)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(40)
    p.paragraph_format.space_after = Pt(8)
    add_run(p, "LAPORAN ANALISIS\nSEMENTARA", bold=True, size=27, color=GREEN)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    add_run(p, "Perbandingan Timbangan Digital Kebun dan Timbangan RAM", bold=True, size=16, color=CHARCOAL)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(24)
    add_run(
        p,
        "Evaluasi selisih berat DO, faktor waktu, tekanan tumpukan, dan kontrol proses",
        size=11.5,
        color=MID_GREY,
    )

    cover_rule = doc.add_table(rows=1, cols=1)
    cover_rule.alignment = WD_TABLE_ALIGNMENT.LEFT
    write_cell(
        cover_rule.cell(0, 0),
        "STATUS: DRAFT UNTUK VALIDASI LAPANGAN",
        bold=True,
        size=9.5,
        color=GREEN,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        fill=PALE_GREEN,
    )
    set_cell_border(
        cover_rule.cell(0, 0),
        top={"val": "single", "sz": 8, "color": GREEN},
        bottom={"val": "single", "sz": 8, "color": GREEN},
        left={"val": "single", "sz": 8, "color": GREEN},
        right={"val": "single", "sz": 8, "color": GREEN},
    )

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(96)
    add_run(p, "Periode data", bold=True, size=8.5, color=MID_GREY)
    add_run(p, "\nJuli-Agustus 2026", bold=True, size=11, color=CHARCOAL)
    p = doc.add_paragraph()
    add_run(p, "Tanggal penyusunan", bold=True, size=8.5, color=MID_GREY)
    add_run(p, "\n26 Agustus 2026", bold=True, size=11, color=CHARCOAL)
    p = doc.add_paragraph()
    add_run(p, "Penggunaan", bold=True, size=8.5, color=MID_GREY)
    add_run(p, "\nInternal PT Julang Plantations", bold=True, size=11, color=CHARCOAL)
    doc.add_page_break()

    # Document control
    add_section_title(doc, "00", "Status dan Dasar Penyusunan")
    control = doc.add_table(rows=4, cols=2)
    control.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_fixed_layout(control)
    control_rows = [
        ("Status dokumen", "Analisis sementara - belum menjadi keputusan final atau dasar koreksi transaksi"),
        ("Sumber utama", "Workbook 'perbandingan susut ram dan digital.xlsx' dan tabel empat perbandingan DO"),
        ("Ruang lingkup", "Perbandingan berat, proses loading, waktu menuju RAM, tekanan tumpukan, dan hipotesis operasional"),
        ("Batasan utama", "Data persentase susut tidak menyertakan kuantitas dasar/formula; pasangan digital-RAM baru empat DO"),
    ]
    for i, (label, value) in enumerate(control_rows):
        write_cell(control.cell(i, 0), label, bold=True, size=8.7, color=GREEN, fill=PALE_GREEN)
        write_cell(control.cell(i, 1), value, size=8.7, fill=WHITE)
    set_table_borders(control, LIGHT_GREY, 5)
    add_callout(
        doc,
        "Catatan konsistensi data",
        "Workbook menandai seluruh baris Juli sebagai 'RAM'. Laporan ini mengasumsikan maksud penjelasan awal adalah Juli belum menggunakan timbangan digital. Apabila label workbook tidak benar, ringkasan bulanan perlu dihitung ulang.",
        fill=PALE_GOLD,
        accent=GOLD,
    )

    # Executive summary
    add_section_title(doc, "01", "Ringkasan Eksekutif")
    add_kpi_cards(doc)
    add_body_paragraph(
        doc,
        "Empat DO yang mempunyai pasangan pengukuran menunjukkan net digital 40.668 kg dan net RAM 35.060 kg. Selisih kumulatifnya 5.608 kg atau 13,79% dari net digital. Pada seluruh pasangan, angka digital lebih tinggi daripada RAM.",
    )
    add_callout(
        doc,
        "Kesimpulan sementara",
        "Data belum membuktikan bahwa salah satu timbangan salah. Digital menimbang lot lebih awal, sedangkan RAM menimbang truk setelah proses loading, penambahan waktu sekitar 0,5-1 jam, dan perjalanan. Kedua sistem kemungkinan mengukur material pada kondisi fisik dan titik waktu yang berbeda.",
    )
    add_bullet(doc, "Tekanan tumpukan dapat berkontribusi apabila tekanan tersebut mengeluarkan air/serum yang kemudian benar-benar meninggalkan bak truk.")
    add_bullet(doc, "Tekanan tumpukan hanya dapat menurunkan berat muatan apabila air/serum atau material benar-benar keluar dari bak truk.")
    add_bullet(doc, "Selisih 790-2.103 kg per truk terlalu besar untuk langsung dianggap sebagai susut normal tanpa pengukuran runoff dan rekonsiliasi lot.")
    add_bullet(doc, "Prioritas awal adalah membuktikan kesetaraan basis gross, kelengkapan lot yang termuat, tare aktual, dan kontribusi air.")

    # Process comparison
    add_section_title(doc, "02", "Perubahan Proses Penimbangan")
    add_process_flow(doc)
    add_body_paragraph(
        doc,
        "Perubahan proses menciptakan titik ukur baru sebelum muatan selesai mengalami tekanan dan perjalanan. Karena itu, perbedaan digital-RAM tidak otomatis sama dengan kesalahan alat. Selisih dapat mencampurkan perubahan kadar air, ketidaksesuaian lot yang termuat, perubahan tare, kehilangan fisik, serta bias kedua alat.",
    )
    add_callout(
        doc,
        "Pertanyaan kunci yang harus dikonfirmasi",
        "Apakah 'gross digital kebun' merupakan pembacaan gross langsung dari alat, atau angka rekonstruksi dari jumlah net seluruh lot ditambah tare kendaraan? Jika merupakan rekonstruksi, gross digital dan gross RAM bukan dua pembacaan yang setara.",
        fill=PALE_GOLD,
        accent=GOLD,
    )

    # Paired analysis
    add_section_title(doc, "03", "Analisis Empat DO Berpasangan")
    add_paired_table(doc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    run = p.add_run()
    run.add_picture(str(CHART_PATH), width=Cm(16.5))
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(6)
    add_run(p2, "Gambar 1. Net digital selalu lebih tinggi pada empat pasangan DO.", size=8, color=MID_GREY, italic=True)

    add_body_paragraph(doc, "Temuan dari pasangan data", bold_lead="Temuan dari pasangan data")
    add_bullet(doc, "Karena tare pada setiap pasangan identik, selisih net pada tabel persis sama dengan selisih gross. Rumus net = gross - tare bukan sumber gap matematis.")
    add_bullet(doc, "Dua DO hari yang sama tetap turun total 2.392 kg atau 11,49% dari net digital. Waktu menginap bukan satu-satunya penyebab.")
    add_bullet(doc, "Dua DO hari berikutnya turun total 3.216 kg atau 16,20%, tetapi seluruhnya menggunakan BK 8035 WS; efek waktu dan kendaraan belum dapat dipisahkan.")
    add_bullet(doc, "Gembung turun 3.705 kg atau 18,84%, sedangkan Sebayur 1.903 kg atau 9,06%. Ini indikasi awal perbedaan kondisi material, bukan bukti final.")
    add_bullet(doc, "Gross digital 14.633 kg muncul pada dua DO berbeda untuk BK 8035 WS dan patut ditelusuri sumber angkanya.")

    # Workbook summary
    add_section_title(doc, "04", "Ringkasan Persentase Susut")
    add_stats_table(doc, source_rows)
    add_body_paragraph(
        doc,
        "Arah data konsisten dengan net digital yang lebih tinggi: rata-rata angka Susut Gabungan pada tanggal berlabel Digital lebih rendah. Namun, perbandingan ini belum dapat dianggap sebagai dampak kausal timbangan karena jumlah sampel kecil, pemilihan metode tidak acak, dan volume setiap DO tidak tersedia.",
    )
    add_callout(
        doc,
        "Cara membaca tabel",
        "Rata-rata di atas adalah rata-rata sederhana antar tanggal, bukan rata-rata berbobot tonase. Angka nol pada kolom produksi diperlakukan sebagai tidak ada nilai untuk produksi tersebut karena Susut Gabungan mengikuti kolom yang berisi nilai.",
        fill=VERY_LIGHT_GREY,
        accent=MID_GREY,
    )

    # Cause matrix in landscape
    landscape = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_section(landscape, landscape=True)
    add_section_title(doc, "05", "Matriks Penyebab yang Mungkin")
    add_evidence_matrix(doc)
    add_callout(
        doc,
        "Interpretasi",
        "Urutan prioritas menunjukkan pemeriksaan yang paling bernilai, bukan tingkat kepastian statistik. Beberapa penyebab dapat terjadi bersamaan dan saling memperbesar selisih.",
        fill=PALE_GOLD,
        accent=GOLD,
    )

    portrait = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_section(portrait, landscape=False)

    # Compression deep dive
    add_section_title(doc, "06", "Analisis Khusus: Tekanan Tumpukan Saat Loading")
    add_callout(
        doc,
        "Definisi operasional",
        "Tekanan tumpukan adalah gaya per satuan luas yang diterima lot akibat berat muatan di atasnya. Tekanan terbesar diperkirakan terjadi pada lot di lapisan bawah. Hipotesis yang diuji adalah tekanan tersebut mendorong keluarnya air/serum sebelum kendaraan ditimbang di RAM.",
        fill=PALE_GOLD,
        accent=GOLD,
    )
    add_callout(
        doc,
        "Prinsip massa",
        "Tekanan sendiri tidak menghilangkan massa. Berat hanya berkurang bila air/serum atau material benar-benar keluar dari bak. Udara yang terdesak tidak dapat menjelaskan kehilangan ratusan hingga ribuan kilogram.",
    )
    add_body_paragraph(
        doc,
        "Pada lantai timbangan digital, lot relatif berdiri sendiri dan mengalami tekanan kecil. Setelah dimuat, lapisan bawah menanggung berat seluruh tumpukan. Getaran, kemiringan jalan, dan waktu menunggu dapat mempercepat keluarnya air yang sebelumnya berada di permukaan, celah, atau bagian dalam lump. Mekanisme ini masuk akal, khususnya untuk material basah, tetapi besarnya harus dibuktikan melalui neraca massa.",
    )
    add_compression_volume_table(doc)
    add_body_paragraph(
        doc,
        "Jumlah ekuivalen cairan tersebut seharusnya menghasilkan runoff yang terlihat. Apabila cairan yang benar-benar keluar jauh lebih kecil, efek tekanan hanya menjelaskan sebagian selisih dan audit harus kembali pada rekonsiliasi lot, tare, dan alat timbang.",
    )
    add_body_paragraph(
        doc,
        "Literatur pendukung", bold_lead="Literatur pendukung"
    )
    add_body_paragraph(
        doc,
        "Penelitian Rubber Research Institute of India melaporkan bahwa cup lump dapat menahan kelembapan, termasuk air yang terperangkap di dalam material, dan jumlahnya dipengaruhi ukuran lump. Pada fresh cup lump, kadar air yang dilaporkan berada pada 47%-53% dari berat basah. Angka ini menunjukkan adanya reservoir air yang potensial, tetapi tidak membuktikan bahwa seluruhnya dapat keluar akibat loading dalam waktu 0,5-1 jam.",
    )

    # Validation plan
    add_section_title(doc, "07", "Rencana Validasi Lapangan")
    add_validation_plan(doc)
    add_body_paragraph(doc, "Aturan keputusan", bold_lead="Aturan keputusan")
    add_decision_rules(doc)

    # Interim recommendations
    add_section_title(doc, "08", "Rekomendasi Operasional Sementara")
    recommendations = [
        "Gunakan net digital sebagai catatan berat lot pada saat dispatch/loading dan net RAM sebagai berat truk setelah proses loading/perjalanan. Jangan menyebut salah satunya salah sebelum uji terkontrol selesai.",
        "Catat selisih digital-RAM sebagai susut transit sementara dengan status belum terklasifikasi; jangan langsung menimpa histori berat sumber lot.",
        "Wajibkan checklist lot tertimbang versus lot termuat, waktu proses, segel kendaraan, foto kondisi bak, serta tiket RAM asli.",
        "Gunakan tare kosong aktual pada perjalanan uji dan dokumentasikan pengemudi, bahan bakar, terpal, alat, serta kondisi basah kendaraan.",
        "Pisahkan analisis berdasarkan produksi, kendaraan, jeda waktu, cuaca, usia material, dan operator. Hindari menyimpulkan dari rata-rata gabungan saja.",
        "Tentukan satu basis komersial resmi setelah validasi: berat saat dispatch, berat penerimaan/RAM, atau skema rekonsiliasi dengan klasifikasi susut yang disepakati.",
    ]
    for recommendation in recommendations:
        add_number(doc, recommendation)
    add_callout(
        doc,
        "Posisi yang disarankan saat ini",
        "Pertahankan kedua angka beserta waktu dan sumbernya. Perbedaan harus menjadi transaksi rekonsiliasi yang dapat ditelusuri, bukan koreksi diam-diam terhadap histori lot atau bukti tunggal bahwa timbangan digital bermasalah.",
    )

    # Conclusion
    add_section_title(doc, "09", "Kesimpulan")
    add_body_paragraph(
        doc,
        "Data sementara menunjukkan pola sistematis: net digital selalu lebih tinggi daripada net RAM, dengan gap kumulatif 5.608 kg atau 13,79%. Perbedaan titik waktu dan kondisi material membuat tekanan tumpukan yang memicu drainase menjadi hipotesis yang masuk akal. Namun, ekuivalen kehilangan 790-2.103 liter per truk cukup besar sehingga harus terlihat dan dapat diukur jika benar-benar berasal dari cairan.",
    )
    add_body_paragraph(
        doc,
        "Belum ada dasar yang cukup untuk menetapkan digital atau RAM sebagai sumber kesalahan tunggal. Kesimpulan final memerlukan audit empat DO, uji tekanan tumpukan dengan penampungan cairan, perjalanan terkontrol, tare aktual, dan verifikasi kedua alat. Fokus manajemen sebaiknya bukan memilih alat mana yang dipercaya terlebih dahulu, melainkan memastikan kedua pengukuran dapat direkonsiliasi melalui bukti lot, waktu, cairan, kendaraan, dan tiket timbang.",
    )

    # Appendix
    doc.add_page_break()
    add_section_title(doc, "A", "Lampiran Data Workbook")
    add_body_paragraph(
        doc,
        "Tabel berikut menyajikan kembali seluruh isi sheet 'Ekstrak Susut'. Tanda '-' menggantikan nilai nol pada produksi yang tidak memiliki angka susut pada tanggal tersebut.",
    )
    add_source_appendix(doc, source_rows)

    add_section_title(doc, "B", "Kolom Data Minimum untuk Analisis Lanjutan")
    add_recommended_data_fields(doc)

    add_section_title(doc, "C", "Referensi")
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.first_line_indent = Cm(-0.5)
    add_run(
        p,
        "Nair, D. B., Jacob, J., & Nair, N. R. (2012). A simple method for rapid determination of residual water content in rubber cup lumps. Journal of Plantation Crops, 40(1), 35-39. ",
        size=9,
        color=CHARCOAL,
    )
    add_run(
        p,
        "https://updatepublishing.com/journal/index.php/JPC/article/view/7575",
        size=9,
        color=GREEN_2,
    )
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.first_line_indent = Cm(-0.5)
    add_run(
        p,
        "PT Julang Plantations. (2026). Perbandingan susut RAM dan digital [Workbook internal dan tabel perbandingan empat DO].",
        size=9,
        color=CHARCOAL,
    )

    # Core properties
    doc.core_properties.title = "Laporan Analisis Sementara Timbangan Digital vs RAM"
    doc.core_properties.subject = "Analisis selisih berat DO dan faktor tekanan tumpukan"
    doc.core_properties.author = "PT Julang Plantations"
    doc.core_properties.keywords = "timbangan digital, timbangan RAM, DO, karet, susut, tekanan tumpukan"
    doc.core_properties.comments = "Disusun berdasarkan data yang diberikan; status sementara untuk validasi lapangan."
    doc.core_properties.created = datetime(2026, 8, 26)

    doc.save(OUTPUT_DOCX)
    return OUTPUT_DOCX


if __name__ == "__main__":
    output = create_document()
    print(output)
